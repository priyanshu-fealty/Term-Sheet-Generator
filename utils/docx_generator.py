"""Document Generation Utilities

This module provides utilities for generating DOCX documents from text content.
"""

import os
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx_from_text(text_content: str, output_path: str, company_name: Optional[str] = None) -> str:
    """Create a DOCX document from text content.

    Args:
        text_content: The text content to convert to DOCX
        output_path: The path to save the DOCX document
        company_name: Optional company name for the header

    Returns:
        The path to the created DOCX document
    """
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse the text content and add to document with formatting
    lines = text_content.split('\n')
    current_section = None
    in_list = False
    list_indent = 0
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines
        if not line:
            doc.add_paragraph()
            continue
        
        # Check if line is a header (starts with # or ##)
        if line.startswith('# '):
            # Main header
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(14)
            current_section = None
            in_list = False
        elif line.startswith('## '):
            # Section header
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(12)
            current_section = line[3:]
            in_list = False
        elif line.startswith('### '):
            # Subsection header
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.italic = True
            current_section = line[4:]
            in_list = False
        elif line.startswith('- ') or line.startswith('* '):
            # List item
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(line[2:])
            in_list = True
        elif line.startswith('  - ') or line.startswith('  * '):
            # Indented list item
            p = doc.add_paragraph()
            p.style = 'List Bullet 2'
            p.add_run(line[4:])
            in_list = True
        elif line.startswith('(') and line[1].isdigit() and line.startswith(') ', 2):
            # Numbered list detected like (1) Item
            p = doc.add_paragraph()
            p.style = 'List Number'
            p.add_run(line[4:])
            in_list = True
        elif line.startswith('---'):
            # Horizontal line
            doc.add_paragraph('_' * 50)
            in_list = False
        else:
            # Regular paragraph
            if in_list and line.startswith('  '):
                # Continuation of list item
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.add_run('\n' + line.strip())
            else:
                p = doc.add_paragraph()
                if '**' in line:
                    # Handle bold text
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Odd parts are inside ** markers
                            p.add_run(part).bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(line)
                in_list = False
    
    # Ensure the directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save the document
    doc.save(output_path)
    
    return output_path


def text_to_docx(text_content: str, output_filename: str = "term_sheet.docx") -> str:
    """Convert text content to a DOCX document.

    Args:
        text_content: The text content to convert
        output_filename: The filename for the output DOCX

    Returns:
        The path to the created DOCX document
    """
    # Extract company name if present
    company_name = None
    lines = text_content.split('\n')
    for line in lines[:10]:  # Check first 10 lines
        if 'COMPANY NAME' in line or 'INC.' in line:
            company_name = line.strip()
            break
    
    # Create the document
    return create_docx_from_text(text_content, output_filename, company_name)


# Example usage
if __name__ == "__main__":
    sample_text = """
    # TERM SHEET FOR SERIES A PREFERRED STOCK FINANCING
    # EXAMPLE COMPANY, INC.
    
    ## FINANCING SUMMARY
    
    **Type of Security:** Series A Preferred Stock
    **Investment Amount:** $5,000,000
    **Pre-Money Valuation:** $20,000,000
    **Discount:** 20%
    
    ## TERMS OF SERIES A PREFERRED STOCK
    
    **Liquidation Preference:** 1x non-participating
    """
    
    output_path = text_to_docx(sample_text)
    print(f"Document created at: {output_path}")